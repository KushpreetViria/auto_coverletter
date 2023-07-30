from docx import Document
from datetime import datetime
from docx2pdf import convert
import PyPDF4
import os


class DocxEditor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = Document(self.file_path)

    @staticmethod
    def get_number_of_pdf_pages(document: str) -> str:
        file = open(document, "rb")
        pdfReader = PyPDF4.PdfFileReader(file)
        return len(pdfReader.pages)

    @staticmethod
    def save_as_pdf(document: str) -> str:
        convert(document)  # Convert to PDF using docx2pdf
        return ".pdf".join(document.rsplit(".docx", 1))

    def populate_data(self, *, street:str, city:str, company_full:str, company_short:str, position:str):
        """
        Replaces occureses of <> in the document with expected values, keep formating
        """
        replacer_map = {
            "<DATE>": datetime.today().strftime("%B %d, %Y"),
            "<STREET_ADDRESS>": street,
            "<CITY_PROVINCE>": city,
            "<COMPANY_NAME_FULL>": company_full,
            "<COMPANY_NAME_SHORT>": company_short,
            "<POSITION_NAME>": position,
        }

        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                for k in replacer_map.keys():
                    replacement = replacer_map.get(k) if replacer_map.get(k) != None else ""
                    run.text = run.text.replace(k, replacement)
        
        # if the top header is missing info, delete the empty paragraphs this leaves behind
        if (not street or not city or not company_full):
            # We skip most of the empty paragraphs, only want to remove the first couple in the header
            skip = 10 - sum(1 if value is None else 0 for value in replacer_map.values())
            for i in reversed(range(len(self.document.paragraphs))):
                paragraph = self.document.paragraphs[i]
                if len(paragraph.text.strip()) == 0:
                    if (skip):
                        skip -= 1
                        continue
                    else:
                        self.document._element.body.remove(paragraph._element)

    def save(self, output_path:str=None) -> str:
        """
        Save the modified document to the specified output_path.
        Returns document name
        """
        name = "modified_document.docx"
        if (output_path):
            folders = output_path.rsplit("/",2)
            if len(folders) > 1 and not os.path.exists(folders[0]):
                os.makedirs(folders[0])
            name = output_path + ".docx"
        self.document.save(name)
        return name
